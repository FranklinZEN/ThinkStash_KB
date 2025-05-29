import asyncio
import os
import sys
from typing import Tuple, List, cast

# Add project root to sys.path to allow anaiservice imports
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# PROJECT_ROOT is the parent directory of 'aiservice' package, which is parent of 'scripts'
PROJECT_ROOT = os.path.dirname(os.path.dirname(SCRIPT_DIR)) 
sys.path.insert(0, PROJECT_ROOT)

from aiservice.app.services.acquisition.file_service import FileAcquisitionService, FileAcquisitionServiceInput
from aiservice.app.services.base import ServiceResult
from aiservice.app.models.pipeline_models import PreliminaryBlock, DocumentMetadata, RawImageInput

# --- Configuration: Paths to your test files ---
# Create these files in a 'test_data' directory or similar next to your script, or provide absolute paths.
# Ensure they contain a mix of elements like headings, lists, images.
TEST_FILES_BASE_PATH = os.path.join(SCRIPT_DIR, "test_data_file_service") # Example: create a test_data_file_service subdir

# Make sure to create these files or update paths for your testing
TEST_TXT_PATH = os.path.join(TEST_FILES_BASE_PATH, "sample.txt")
TEST_MD_PATH = os.path.join(TEST_FILES_BASE_PATH, "sample.md") 
# For MD, include local images (e.g. in a subfolder like 'md_images/local_image.png') and remote image URLs
TEST_DOCX_PATH = os.path.join(TEST_FILES_BASE_PATH, "sample.docx")


async def run_file_acquisition_test(file_path: str, content_type: str, test_job_id: str):
    print(f"\n--- Testing FileAcquisitionService with: {file_path} (Type: {content_type}) ---")
    
    if not os.path.exists(file_path):
        print(f"Test file not found: {file_path}. Please create it or update the path.")
        return

    file_service = FileAcquisitionService()
    service_input = FileAcquisitionServiceInput(
        file_path=file_path,
        source_content_type=content_type,
        processing_level="full_content",
        job_id=test_job_id
    )

    service_result: ServiceResult[Tuple[List[PreliminaryBlock], DocumentMetadata, List[RawImageInput]]] = await file_service.execute(service_input)

    if service_result.success and service_result.data:
        preliminary_blocks, doc_metadata, raw_images = service_result.data

        print("\n--- DocumentMetadata ---")
        print(f"  Document ID: {doc_metadata.document_id}")
        print(f"  Title: {doc_metadata.title}")
        print(f"  Source: {doc_metadata.source_identifier}")
        print(f"  Source Type: {doc_metadata.source_type}")
        print(f"  Author: {doc_metadata.author}")
        print(f"  Subject: {doc_metadata.subject}")
        print(f"  Keywords: {doc_metadata.keywords}")
        print(f"  Creation Date: {doc_metadata.creation_date}")
        print(f"  Modification Date: {doc_metadata.modification_date}")
        print(f"  Extracted At: {doc_metadata.extracted_at}")
        if hasattr(doc_metadata, 'total_pages') and doc_metadata.total_pages is not None:
            print(f"  Total Pages: {doc_metadata.total_pages}") # Relevant for paginated sources, maybe not files

        print(f"\n--- PreliminaryBlocks ({len(preliminary_blocks)}) ---")
        text_block_count = 0
        heading_block_count = 0
        list_item_block_count = 0
        image_placeholder_count = 0
        code_snippet_count = 0

        for i, block in enumerate(preliminary_blocks):
            print(f"  Block {i}: ID={block.block_id}, Type='{block.type}', Order={block.order}")
            if block.type == "text":
                text_block_count += 1
                print(f"    Text: '{block.text_content[:100]}...'")
            elif block.type == "heading":
                heading_block_count +=1
                print(f"    Heading (L{block.heading_level}): '{block.text_content[:100]}...'")
            elif block.type == "list_item":
                list_item_block_count +=1
                print(f"    List Item (L{block.list_level}, Ordered={block.list_ordered}): '{block.text_content[:100]}...'")
            elif block.type == "image_placeholder":
                image_placeholder_count +=1
                print(f"    Image ID Ref: {block.image_id_ref}")
            elif block.type == "code_snippet":
                code_snippet_count += 1
                print(f"    Code ({block.code_language or 'N/A'}): \n{block.code_content[:200]}...")
        
        print(f"  Total Text Blocks: {text_block_count}")
        print(f"  Total Heading Blocks: {heading_block_count}")
        print(f"  Total List Item Blocks: {list_item_block_count}")
        print(f"  Total Image Placeholders: {image_placeholder_count}")
        print(f"  Total Code Snippet Blocks: {code_snippet_count}")

        is_order_sequential = all(preliminary_blocks[i].order == i for i in range(len(preliminary_blocks)))
        print(f"  PreliminaryBlock.order assignment is sequential and correct: {is_order_sequential}")

        print(f"\n--- RawImageInputs ({len(raw_images)}) ---")
        for i, img_input in enumerate(raw_images):
            print(f"  Image {i}: ID={img_input.image_id}")
            print(f"    Source URL: {img_input.source_url}")
            print(f"    Original Filename: {img_input.original_filename}")
            print(f"    MIME Type: {img_input.mime_type}")
            print(f"    Alt Text: {img_input.alt_text}")
            print(f"    Has Bytes: {img_input.image_bytes is not None} (Length: {len(img_input.image_bytes) if img_input.image_bytes else 'N/A'})")
            print(f"    GCS Path Components: job_id='{img_input.job_id_for_gcs_path}', source_type='{img_input.source_type_for_gcs_path}', original_source='{img_input.original_source_identifier_for_gcs_path}'")

        # Verification
        print("\n--- Verifications ---")
        if image_placeholder_count > 0 or len(raw_images) > 0:
            placeholder_ids = {pb.image_id_ref for pb in preliminary_blocks if pb.type == "image_placeholder"}
            raw_image_ids = {ri.image_id for ri in raw_images}
            if placeholder_ids == raw_image_ids:
                print("  All image_id_refs in image_placeholders successfully match RawImageInput IDs.")
            else:
                print(f"  ERROR: Mismatch in image IDs!")
                print(f"    Placeholder Refs: {placeholder_ids}")
                print(f"    RawImage IDs: {raw_image_ids}")
        else:
            print("  No images or placeholders to verify.")

    else:
        print(f"Service call failed for {file_path}.")
        if service_result.error_message:
            print(f"  Error: {service_result.error_message}")
        if service_result.error_details:
            print(f"  Details: {service_result.error_details}")

    print("\n--- Test Complete ---")

async def main():
    # Create dummy files if they don't exist for a basic run, user should replace with real test data
    os.makedirs(TEST_FILES_BASE_PATH, exist_ok=True)
    if not os.path.exists(TEST_TXT_PATH):
        with open(TEST_TXT_PATH, 'w') as f:
            f.write("This is paragraph one.\n\nThis is paragraph two with a *star*.")
        print(f"Created dummy {TEST_TXT_PATH}")
    
    if not os.path.exists(TEST_MD_PATH):
        os.makedirs(os.path.join(TEST_FILES_BASE_PATH, "md_images"), exist_ok=True)
        # Create a dummy local image for markdown test
        dummy_md_image_path = os.path.join(TEST_FILES_BASE_PATH, "md_images", "local_img.png")
        if not os.path.exists(dummy_md_image_path):
            try:
                from PIL import Image, ImageDraw
                img = Image.new('RGB', (60, 30), color = 'red')
                d = ImageDraw.Draw(img)
                d.text((10,10), "IMG", fill=(255,255,0))
                img.save(dummy_md_image_path, 'PNG')
                print(f"Created dummy image {dummy_md_image_path}")
            except ImportError:
                print("Pillow (PIL) not installed, cannot create dummy image. Please create one manually or install Pillow.")
                with open(dummy_md_image_path, 'wb') as fimg: fimg.write(b'dummyimagebytes') # placeholder

        with open(TEST_MD_PATH, 'w') as f:
            f.write("# Main Heading\n\nThis is some text.\n\n* Item 1\n* Item 2\n\n1. Ordered 1\n2. Ordered 2\n\n```python\nprint(\"Hello\")\n```\n\n![Alt text for local image](md_images/local_img.png)\n\n![Alt text for remote image](https://www.google.com/images/branding/googlelogo/1x/googlelogo_color_272x92dp.png)")
        print(f"Created dummy {TEST_MD_PATH}")

    # For DOCX, it's harder to create a dummy. User should provide one.
    if not os.path.exists(TEST_DOCX_PATH):
        print(f"WARNING: Test file {TEST_DOCX_PATH} not found. Please create a sample DOCX file for testing.")
        # Basic docx creation if python-docx is available, but will be very simple
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Test Heading', level=1)
            doc.add_paragraph('This is a test paragraph.')
            doc.add_paragraph('Item 1', style='ListBullet')
            # Not adding image to dummy docx as it's more involved.
            doc.save(TEST_DOCX_PATH)
            print(f"Created basic dummy {TEST_DOCX_PATH}")
        except ImportError:
            print("python-docx not installed, cannot create dummy DOCX. Please create one manually.")
        except Exception as e_docx_create:
            print(f"Could not create dummy DOCX: {e_docx_create}")

    # Run tests
    if os.path.exists(TEST_TXT_PATH):
        await run_file_acquisition_test(TEST_TXT_PATH, "txt", "test_txt_job_001")
    if os.path.exists(TEST_MD_PATH):
        await run_file_acquisition_test(TEST_MD_PATH, "md", "test_md_job_001")
    if os.path.exists(TEST_DOCX_PATH):
        await run_file_acquisition_test(TEST_DOCX_PATH, "docx", "test_docx_job_001")

if __name__ == "__main__":
    asyncio.run(main()) 